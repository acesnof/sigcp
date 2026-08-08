from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "_util" / "Manual_SIGCP_Utilizacao_e_Administracao.docx"
TEAL = "0B5962"
TEAL_DARK = "073F46"
RED = "B51618"
PALE = "EAF3F4"
GREY = "60777A"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    tc_pr.append(fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def add_field(paragraph, instruction, placeholder="Atualizar campo"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    keep_with_next(paragraph)
    return paragraph


def body(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def bullets(doc, items):
    for item in items:
        body(doc, item, "List Bullet")


def steps(doc, items):
    for item in items:
        body(doc, item, "List Number")


def note(doc, title, text, color=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, color)
    set_cell_margins(cell, 140, 170, 140, 170)
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL_DARK)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def screenshot(doc, code, caption, guidance):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, "F4F7F7")
    set_cell_margins(cell, 260, 180, 260, 180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ INSERIR SCREENSHOT {code} ]")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(RED)
    p2 = cell.add_paragraph(caption)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    p3 = cell.add_paragraph(guidance)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].italic = True
    p3.runs[0].font.size = Pt(9)
    body(doc, f"Figura {code} — {caption}").alignment = WD_ALIGN_PARAGRAPH.CENTER


def matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, TEAL)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_margins(cells[index])
            cells[index].text = str(value)
            if row_index % 2:
                shade(cells[index], "F5F8F8")
            if widths:
                cells[index].width = Cm(widths[index])
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string("163438")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 34, TEAL_DARK), ("Heading 1", 21, TEAL_DARK), ("Heading 2", 15, TEAL), ("Heading 3", 11, RED)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    header = section.header.paragraphs[0]
    header.text = "SIGCP  |  MANUAL DE UTILIZAÇÃO E ADMINISTRAÇÃO"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("USO INTERNO  •  Página ")
    add_field(footer, "PAGE", "1")
    footer.add_run(" de ")
    add_field(footer, "NUMPAGES", "1")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GREY)
    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def build():
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SIGCP")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(42)
    r.font.color.rgb = RGBColor.from_string(TEAL_DARK)
    p = doc.add_paragraph("Sistema Integrado de Gestão do Contingente Português")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph("MANUAL DE UTILIZAÇÃO E ADMINISTRAÇÃO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(19)
    p.runs[0].font.color.rgb = RGBColor.from_string(RED)
    body(doc, "Guia funcional, operacional e técnico").alignment = WD_ALIGN_PARAGRAPH.CENTER
    body(doc, "Edição 1.0  |  Agosto de 2026  |  Uso interno").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_before = Pt(90)
    note(doc, "Documento editável", "Os blocos assinalados indicam exatamente onde inserir os screenshots finais. Atualize o índice no Word depois de inserir as imagens.")
    page_break(doc)

    heading(doc, "Controlo do documento", 1)
    matrix(doc, ["Campo", "Informação"], [
        ("Título", "Manual de Utilização e Administração do SIGCP"),
        ("Versão", "1.0"), ("Data", "Agosto de 2026"),
        ("Classificação", "Uso interno"),
        ("Aplicação", "SIGCP.exe — aplicação local com interface no browser"),
        ("Responsável pelo documento", "[PREENCHER]"),
        ("Aprovado por", "[PREENCHER]"),
    ])
    heading(doc, "Histórico de revisões", 2)
    matrix(doc, ["Versão", "Data", "Alteração", "Autor"], [("1.0", "01-08-2026", "Emissão inicial do manual completo", "[PREENCHER]")])
    heading(doc, "Como finalizar este manual", 2)
    steps(doc, [
        "Substituir cada caixa ‘INSERIR SCREENSHOT’ pela imagem indicada.",
        "No Word, clicar com o botão direito no índice e escolher Atualizar campo > Atualizar índice inteiro.",
        "Rever nomes, contactos, caminhos de rede e política local de cópias de segurança.",
        "Guardar a versão editável em DOCX e emitir uma cópia PDF controlada.",
    ])
    page_break(doc)
    heading(doc, "Índice", 1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Clique com o botão direito e selecione Atualizar campo")
    page_break(doc)

    heading(doc, "1. Visão geral", 1)
    body(doc, "O SIGCP centraliza o planeamento mensal, os Welfares individuais, a gestão de férias, o pessoal, os documentos operacionais e a administração da base de dados. A aplicação é executada no computador e disponibiliza a interface através do browser, num endereço local protegido por autenticação.")
    heading(doc, "1.1 Princípios de funcionamento", 2)
    bullets(doc, [
        "Uma única base de dados SQLite contém a informação operacional.",
        "O acesso a páginas e ações depende dos Tipos de Acesso e das funções SNR/Responsável Welfare.",
        "As alterações relevantes ficam registadas na Auditoria.",
        "O login desencadeia uma cópia de segurança automática; uma falha de backup é comunicada sem impedir o acesso.",
        "A ordenação de pessoas coloca primeiro quem está na missão; em cada grupo ordena por posto e depois pela data de antiguidade.",
    ])
    screenshot(doc, "01", "Ecrã inicial e menu principal", "Capturar a aplicação após login, com o menu lateral completo visível.")

    heading(doc, "2. Instalação, arranque e encerramento", 1)
    heading(doc, "2.1 Requisitos", 2)
    bullets(doc, ["Windows 10/11.", "Browser moderno: Microsoft Edge, Google Chrome ou Mozilla Firefox.", "Permissão de leitura e escrita na pasta da base de dados e respetiva pasta db_backup.", "Acesso ao caminho de rede, quando a base de dados se encontra num servidor."])
    heading(doc, "2.2 Primeiro arranque", 2)
    steps(doc, ["Executar SIGCP.exe.", "Se não existir uma base de dados configurada, selecionar o ficheiro SQLite indicado pela administração.", "Aguardar a abertura automática do browser.", "Autenticar com o NIM e a palavra-passe atribuída."])
    note(doc, "Base de dados configurada", "A localização escolhida fica guardada no perfil local do Windows. A pasta db_backup é criada ao lado da base de dados, não ao lado do executável.")
    heading(doc, "2.3 Utilização diária", 2)
    body(doc, "Se o SIGCP já estiver em execução, um segundo arranque liga-se à instância existente. Não crie cópias manuais da base de dados enquanto a aplicação está em utilização.")
    heading(doc, "2.4 Encerramento correto", 2)
    steps(doc, ["Usar a opção Encerrar aplicação no menu do utilizador.", "Confirmar o encerramento do servidor local.", "Fechar separadores antigos do SIGCP no browser."])
    screenshot(doc, "02", "Login e opções da sessão", "Capturar o login e, em imagem separada dentro da mesma figura, o menu do perfil/encerramento.")

    heading(doc, "3. Perfis, acessos e responsabilidades", 1)
    body(doc, "Os Tipos de Acesso definem módulos operacionais. As funções SNR e Responsável Welfare são atribuições separadas e apenas os Administradores as podem conceder. Um utilizador pode acumular acessos quando necessário.")
    matrix(doc, ["Perfil/função", "Capacidade principal"], [
        ("Administrador", "Acesso total, pessoal, configurações, auditoria, eliminação de períodos de férias e exportação da base de dados."),
        ("Gestão Welfare Mensal", "Elaboração do plano mensal e edição das refeições no calendário."),
        ("Gestão Ementa", "Consulta da ementa; a edição das refeições exige Gestão Welfare Mensal, Administrador ou Responsável Welfare."),
        ("Gestão Welfare Individual", "Gestão dos Welfares individuais do contingente."),
        ("Pessoal/Gestão Férias", "Acesso funcional à gestão de pessoal/férias, conforme permissões apresentadas."),
        ("Leitura", "Consulta sem alteração dos dados protegidos."),
        ("SNR", "Decisão sobre pedidos de outras pessoas, notificações de gestão e nomeação temporária de substituto."),
        ("Responsável Welfare", "Operações e documentos reservados ao responsável; pode editar refeições no calendário."),
    ], [4.2, 12.0])
    note(doc, "Separação das notificações", "As notificações de ‘As minhas férias’ são pessoais. A bolinha de ‘Gestão de Férias’ pertence ao SNR e contém apenas pedidos de outras pessoas.")

    heading(doc, "4. Navegação e regras comuns", 1)
    bullets(doc, [
        "Planeamento: Calendário Mensal e Welfare Individual.",
        "Férias: As minhas férias e Gestão de Férias.",
        "Pessoal: lista, dados, direitos e substituição SNR.",
        "Sistema: Administração e Auditoria, apenas para Administradores.",
    ])
    body(doc, "As datas e horas são escolhidas em campos separados e alinhados, garantindo compatibilidade entre browsers. Confirme sempre data, hora e minutos antes de guardar.")
    screenshot(doc, "03", "Estrutura de navegação", "Capturar o menu lateral e identificar visualmente as quatro áreas funcionais.")

    heading(doc, "5. Calendário mensal", 1)
    body(doc, "O Calendário Mensal apresenta a atividade de Welfare por dia e refeição. A edição está limitada a Gestão Welfare Mensal, Administrador e Responsável Welfare; os restantes perfis consultam em modo de leitura.")
    heading(doc, "5.1 Registar ou alterar uma refeição", 2)
    steps(doc, ["Escolher o ano e o mês.", "Selecionar o dia/refeição pretendido.", "Preencher os dados disponibilizados no formulário.", "Guardar e confirmar a atualização na grelha."])
    heading(doc, "5.2 Exportação e impressão", 2)
    body(doc, "O botão Exportar PDF gera o planeamento do período apresentado. Antes de emitir, confirme o mês, as marcações, os Day Off e a ementa.")
    screenshot(doc, "04", "Calendário mensal e edição de refeição", "Capturar a grelha e um formulário aberto, incluindo os controlos de mês.")

    heading(doc, "6. Welfare Individual", 1)
    body(doc, "A área individual permite consultar e, quando autorizado, corrigir as marcações por pessoa. Alterações não guardadas são assinaladas e a aplicação pede confirmação antes de abandonar a página.")
    bullets(doc, ["Guardar as alterações pendentes antes de mudar de mês.", "Trancar o mês impede novas alterações; apenas Administradores e Responsáveis Welfare podem gerir o bloqueio.", "Repor marcações recupera os Welfares de origem do período.", "As férias aprovadas refletem-se automaticamente nas refeições abrangidas."])
    heading(doc, "6.1 Documentos e exportações", 2)
    body(doc, "Consoante a permissão, estão disponíveis impressão individual, exportação semanal, Request, Request HOTO, Service Note, Distribuição XFA e documentação de reembolso. Verifique sempre os números de referência indicados nos avisos antes de finalizar o documento.")
    screenshot(doc, "05", "Welfare Individual", "Capturar a grelha com barra de ações, semanas e uma alteração pendente.")
    screenshot(doc, "06", "Opções de impressão e documentos", "Capturar o modal de impressão e o conjunto de ações documentais.")

    heading(doc, "7. As minhas férias", 1)
    body(doc, "Esta área é pessoal: cada utilizador consulta os seus períodos, submete pedidos e acompanha decisões sem misturar as notificações de gestão do SNR.")
    heading(doc, "7.1 Novo pedido", 2)
    steps(doc, ["Clicar em Novo pedido de férias.", "Selecionar a data e hora de partida e de chegada.", "Preencher companhia/voo e observações quando aplicável.", "Ler os avisos de sobreposição ou regras de planeamento.", "Confirmar os avisos e submeter."])
    heading(doc, "7.2 Alteração e cancelamento", 2)
    body(doc, "Uma alteração ou um cancelamento de um período aprovado fica pendente de decisão. Até à aprovação do SNR, o período anteriormente aprovado mantém-se refletido no Welfare Individual.")
    screenshot(doc, "07", "As minhas férias", "Capturar resumo, cartões de períodos e a bolinha de notificações pessoais.")
    screenshot(doc, "08", "Novo pedido de férias", "Capturar o formulário completo, destacando data, hora e minutos alinhados.")

    heading(doc, "8. Gestão de Férias", 1)
    body(doc, "A Gestão de Férias concentra decisões, planeamento, calendário e relatórios. O SNR decide apenas pedidos de outras pessoas; não pode aprovar o seu próprio pedido.")
    heading(doc, "8.1 Lista e filtros", 2)
    bullets(doc, ["Todas: apresenta todos os estados abrangidos.", "Pendentes: apresenta pedidos que requerem decisão.", "Aprovadas: apresenta licenças aprovadas.", "Anuladas: apresenta licenças anuladas.", "Apenas um filtro de estado fica ativo de cada vez; pesquisa e área podem restringir o resultado."])
    body(doc, "As cores dos botões correspondem às cores dos cartões. A impressão usa o filtro ativo e ajusta o título, por exemplo ‘Licenças pendentes de aprovação’.")
    heading(doc, "8.2 Pessoas presentes e histórico", 2)
    body(doc, "Por defeito, Total de Pessoas conta somente quem ainda está na missão e o Excel inclui apenas essas pessoas e os respetivos dados. Ao ativar Incluir Passadas, o total passa a abranger todo o pessoal e o relatório Excel é igualmente alargado a todos.")
    note(doc, "Regra de ordenação", "As pessoas são ordenadas primeiro pelo posto e depois pela data de antiguidade. Quem já saiu da missão aparece no fim, mantendo a mesma ordenação dentro desse grupo.")
    heading(doc, "8.3 Decisões e eliminação", 2)
    body(doc, "O SNR pode aprovar, devolver ou rejeitar pedidos conforme o estado. Os Administradores dispõem do botão Apagar em qualquer fase; esta eliminação é permanente e deve ser usada apenas para corrigir registos indevidos.")
    heading(doc, "8.4 Calendário e impressão mensal", 2)
    body(doc, "No separador Calendário, selecione o mês e use Imprimir para emitir exatamente o mês apresentado, com título e legenda. A ordem das linhas respeita posto e antiguidade.")
    screenshot(doc, "09", "Gestão de Férias — lista e filtros", "Capturar o Total de Pessoas, pesquisa, botões coloridos, Incluir Passadas e Exportar Excel.")
    screenshot(doc, "10", "Detalhe e decisão de um pedido", "Capturar um pedido pendente com os botões de decisão; num perfil Administrador, incluir o botão Apagar.")
    screenshot(doc, "11", "Calendário mensal de férias", "Capturar um mês com legenda, pessoas ordenadas e botão Imprimir.")

    heading(doc, "9. Substituição temporária do SNR", 1)
    body(doc, "O SNR titular pode nomear uma pessoa para assumir temporariamente as funções durante as suas férias. A opção também está disponível aos Administradores.")
    steps(doc, ["Abrir Pessoal e editar o perfil da pessoa escolhida.", "Ativar a checkbox de substituto SNR.", "Indicar data de início e data de fim.", "Guardar.", "Para remover a nomeação, clicar Limpar e guardar."])
    note(doc, "Validade", "As permissões de substituição só estão ativas dentro do intervalo indicado. A pessoa escolhida não pode ser o próprio SNR, o utilizador mestre nem outro SNR titular.")
    screenshot(doc, "12", "Nomeação de substituto SNR", "Capturar a checkbox, datas de início/fim e botão Limpar no perfil da pessoa.")

    heading(doc, "10. Pessoal", 1)
    body(doc, "A lista de pessoal reúne identificação, missão, contacto, área funcional, Posição N.º, antiguidade, acessos e direitos. Quem saiu da missão é colocado no fim da lista.")
    heading(doc, "10.1 Criar ou editar pessoa", 2)
    bullets(doc, ["Preencher NIM, posto, nome e apelido.", "Registar início e fim da missão e data de antiguidade.", "No mesmo grupo, preencher Área funcional e Posição N.º.", "Preencher Telemóvel Serviço quando a pessoa é Responsável Welfare.", "Total de dias Férias (manual) substitui o cálculo automático quando preenchido."])
    heading(doc, "10.2 Acessos e funções", 2)
    body(doc, "Apenas Administradores alteram os Tipos de Acesso e atribuem SNR/Responsável Welfare. Para utilizadores não administradores, o formulário mostra os valores em leitura com essa indicação explícita.")
    screenshot(doc, "13", "Lista de pessoal", "Capturar pessoas presentes e passadas, incluindo ícones de função e ordenação.")
    screenshot(doc, "14", "Adicionar/editar pessoa", "Capturar o modal largo completo, incluindo Área funcional, Posição N.º, direitos, acessos e funções.")

    heading(doc, "11. Administração", 1)
    body(doc, "A Administração é reservada a Administradores e agrega configurações, controlo de utilizadores, regras de planeamento, feriados, Day Off, exportação técnica e Auditoria.")
    heading(doc, "11.1 Configurações e regras", 2)
    body(doc, "Antes de alterar parâmetros, registe o valor anterior e o motivo. Algumas alterações afetam cálculos e relatórios futuros; confirme o ano de calendário e os limites funcionais antes de guardar.")
    heading(doc, "11.2 Feriados e Day Off", 2)
    body(doc, "Os feriados ativos são considerados FS nos cálculos de férias. Os Day Off devem ser mantidos de acordo com a ordem de serviço aplicável.")
    heading(doc, "11.3 Exportar JSON", 2)
    body(doc, "Exportar JSON cria uma cópia lógica das tabelas e registos para diagnóstico ou arquivo. Não substitui o backup SQLite para recuperação operacional.")
    screenshot(doc, "15", "Administração — configurações", "Capturar os separadores da Administração e as principais configurações.")

    heading(doc, "12. Auditoria", 1)
    body(doc, "A Auditoria regista operações concluídas que alteram dados. Credenciais, palavras-passe e tokens são ocultados. Por defeito, são mostradas apenas alterações do mês corrente, reduzindo o volume e acelerando a consulta.")
    heading(doc, "12.1 Pesquisa eficiente", 2)
    steps(doc, ["Abrir Administração > Auditoria.", "Definir datas quando pretende consultar períodos anteriores.", "Usar pesquisa por utilizador, ação, entidade ou conteúdo relevante.", "Aplicar os filtros.", "Abrir o detalhe do registo para comparar dados antes/depois."])
    body(doc, "Os resultados são paginados para manter o desempenho quando existirem muitos registos. Refine datas e texto antes de navegar por páginas adicionais.")
    screenshot(doc, "16", "Auditoria e filtros", "Capturar a tabela sem scroll horizontal, filtros do mês corrente, paginação e um detalhe aberto.")

    heading(doc, "13. Backups automáticos", 1)
    body(doc, "Em cada login, o SIGCP tenta produzir um snapshot consistente da base de dados. O processo usa o mecanismo de backup do SQLite, verifica a integridade da cópia e só depois a disponibiliza na pasta final.")
    heading(doc, "13.1 Localização e retenção", 2)
    bullets(doc, ["Origem: ficheiro SQLite configurado para o SIGCP.", "Destino: pasta db_backup ao lado da base de dados.", "Organização: uma pasta identificada por data e hora para cada execução.", "Retenção automática: as 20 cópias mais recentes.", "Uma falha ao eliminar uma cópia antiga não invalida um backup novo já concluído."])
    heading(doc, "13.2 O que fazer perante um aviso", 2)
    steps(doc, ["Anotar a mensagem apresentada no login.", "Confirmar que existe espaço em disco.", "Confirmar permissões de criação, escrita e leitura no destino.", "Se for rede, confirmar ligação ao servidor e credenciais do Windows.", "Resolver a causa e voltar a autenticar para repetir o backup."])
    note(doc, "Continuidade", "O login não é bloqueado por uma falha de backup, mas o aviso deve ser tratado no próprio dia. Uma aplicação operacional sem backups recentes representa risco de perda de dados.", "FFF1D6")
    screenshot(doc, "17", "Estrutura da pasta db_backup", "Capturar o Explorador do Windows com a base de dados e as pastas de backups por data/hora. Ocultar caminhos sensíveis.")

    heading(doc, "14. Restauro e recuperação", 1)
    note(doc, "Operação administrativa", "O restauro deve ser executado por pessoa autorizada, com todas as instâncias SIGCP encerradas e após preservação da base atual.", "FDE8E8")
    steps(doc, [
        "Informar os utilizadores e encerrar completamente o SIGCP em todos os computadores.",
        "Identificar a última cópia válida na pasta db_backup e confirmar data/hora.",
        "Criar uma cópia de segurança adicional do ficheiro atual, mesmo que esteja danificado.",
        "Copiar o ficheiro SQLite do backup selecionado para a localização configurada, mantendo o nome esperado.",
        "Iniciar o SIGCP num único computador e validar login, pessoal, calendário, férias e auditoria.",
        "Só depois da validação autorizar o acesso geral e documentar o incidente/restauro.",
    ])
    body(doc, "Nunca substitua a base de dados enquanto existir uma aplicação ligada. Em ambiente de rede, confirme também que não existem processos SIGCP.exe ativos noutros postos.")

    heading(doc, "15. Segurança e proteção de dados", 1)
    bullets(doc, ["Atribuir apenas os acessos indispensáveis à função.", "Desativar ou ajustar prontamente perfis de quem termina funções.", "Não partilhar palavras-passe nem enviar a base de dados por canais não autorizados.", "Proteger backups com as mesmas regras da base principal.", "Usar a Auditoria para investigação, controlo e prestação de contas.", "Evitar screenshots com palavras-passe, dados pessoais desnecessários ou caminhos de rede sensíveis."])

    heading(doc, "16. Resolução de problemas", 1)
    matrix(doc, ["Sintoma", "Verificação/ação"], [
        ("O browser não abre", "Confirmar se SIGCP.exe está ativo e abrir http://127.0.0.1:52147 no browser."),
        ("Porta local ocupada", "Fechar instâncias antigas do SIGCP; confirmar no Gestor de Tarefas antes de reiniciar."),
        ("Base de dados não encontrada", "Selecionar novamente o ficheiro correto e confirmar acesso ao caminho de rede."),
        ("Aviso de backup", "Verificar espaço, permissões e conectividade; repetir o login após resolução."),
        ("Não consigo editar refeições", "Confirmar Gestão Welfare Mensal, Administrador ou função Responsável Welfare."),
        ("Não vejo Administração", "A página é exclusiva de Administradores."),
        ("Pedido não aparece ao SNR", "Confirmar se pertence a outra pessoa, estado do pedido e validade da substituição SNR."),
        ("Hora desalinhada ou controlo diferente", "Atualizar o browser; os campos separados de data/hora asseguram compatibilidade com Firefox."),
        ("Excel contém pessoal passado", "Desativar Incluir Passadas antes de exportar."),
        ("Índice do manual desatualizado", "No Word, clicar no índice e escolher Atualizar índice inteiro."),
    ], [5.2, 11.0])

    heading(doc, "17. Rotinas recomendadas", 1)
    matrix(doc, ["Periodicidade", "Rotina"], [
        ("Diária", "Verificar avisos de backup, pedidos pendentes, notificações e alterações do planeamento."),
        ("Semanal", "Confirmar backups recentes, rever substituições SNR e validar dados de pessoas em entrada/saída."),
        ("Mensal", "Rever acessos, trancar o período concluído, exportar relatórios necessários e consultar a Auditoria."),
        ("Antes de rotação", "Validar datas de missão, posto, antiguidade, Área funcional, Posição N.º e direitos de férias."),
        ("Após incidente", "Preservar evidência, consultar Auditoria, validar backups e registar ações de recuperação."),
    ])

    heading(doc, "18. Checklist de passagem de funções", 1)
    bullets(doc, [
        "Localização da base de dados e responsável pelo servidor confirmados.",
        "Último backup válido identificado e procedimento de restauro testado/documentado.",
        "Administradores, SNR e Responsável Welfare atualizados.",
        "Substituições temporárias revistas e datas válidas.",
        "Pessoal presente/passado e respetivas datas de missão validados.",
        "Períodos de férias pendentes tratados.",
        "Meses concluídos trancados e relatórios arquivados.",
        "Auditoria revista para alterações relevantes.",
        "Manual atualizado com screenshots da versão em produção.",
    ])

    heading(doc, "Anexo A — Plano de screenshots", 1)
    rows = [
        ("01", "Ecrã inicial/menu"), ("02", "Login e sessão"), ("03", "Navegação"),
        ("04", "Calendário mensal"), ("05", "Welfare Individual"), ("06", "Documentos"),
        ("07", "As minhas férias"), ("08", "Novo pedido"), ("09", "Gestão de Férias"),
        ("10", "Decisão de pedido"), ("11", "Calendário de férias"), ("12", "Substituto SNR"),
        ("13", "Lista de pessoal"), ("14", "Editar pessoa"), ("15", "Administração"),
        ("16", "Auditoria"), ("17", "Backups"),
    ]
    matrix(doc, ["Código", "Conteúdo", "Inserido/revisto"], [(a, b, "☐") for a, b in rows])

    heading(doc, "Anexo B — Registo local", 1)
    matrix(doc, ["Elemento", "Valor a preencher"], [
        ("Localização da base de dados", "[PREENCHER]"),
        ("Servidor/pasta de rede", "[PREENCHER]"),
        ("Responsável técnico", "[PREENCHER]"),
        ("Administrador funcional", "[PREENCHER]"),
        ("Contacto para incidentes", "[PREENCHER]"),
        ("Local de arquivo dos relatórios", "[PREENCHER]"),
        ("Último teste de restauro", "[PREENCHER]"),
    ])

    doc.core_properties.title = "Manual de Utilização e Administração do SIGCP"
    doc.core_properties.subject = "Guia funcional, operacional e técnico"
    doc.core_properties.author = "SIGCP"
    doc.core_properties.keywords = "SIGCP, welfare, férias, pessoal, auditoria, backups"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
